"""Export Service — generates PDF and DOCX executive reports."""
import io
import os
from datetime import datetime
from typing import List, Dict, Any
from app.models.models import Project, SourceFile


def generate_pdf(project: Project, files: List[SourceFile], phases: List[Dict], executive_summary: str) -> bytes:
    """Generate a PDF executive report using reportlab."""
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor, white, black
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        HRFlowable, PageBreak
    )
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    # Colors
    PRIMARY = HexColor("#6366f1")
    DARK_BG = HexColor("#1e1b4b")
    CRITICAL_COLOR = HexColor("#ef4444")
    HIGH_COLOR = HexColor("#f97316")
    MEDIUM_COLOR = HexColor("#eab308")
    LOW_COLOR = HexColor("#22c55e")
    LIGHT_GRAY = HexColor("#f8fafc")
    MID_GRAY = HexColor("#94a3b8")

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle("Title", parent=styles["Title"],
        fontSize=26, textColor=PRIMARY, spaceAfter=6, fontName="Helvetica-Bold")
    subtitle_style = ParagraphStyle("Subtitle", parent=styles["Normal"],
        fontSize=12, textColor=MID_GRAY, spaceAfter=20)
    h2_style = ParagraphStyle("H2", parent=styles["Heading2"],
        fontSize=16, textColor=DARK_BG, spaceBefore=20, spaceAfter=8, fontName="Helvetica-Bold")
    body_style = ParagraphStyle("Body", parent=styles["Normal"],
        fontSize=10, textColor=HexColor("#334155"), spaceAfter=8, leading=15)
    label_style = ParagraphStyle("Label", parent=styles["Normal"],
        fontSize=9, textColor=MID_GRAY, spaceAfter=2)

    def risk_color(level: str):
        return {"Critical": CRITICAL_COLOR, "High": HIGH_COLOR, "Medium": MEDIUM_COLOR, "Low": LOW_COLOR}.get(level, MID_GRAY)

    story = []

    # ── Cover ───────────────────────────────────────────────────────────────
    story.append(Spacer(1, 0.5 * inch))
    story.append(Paragraph("LegacyLens", title_style))
    story.append(Paragraph("AI-Powered Migration Risk Analysis Report", subtitle_style))
    story.append(HRFlowable(width="100%", thickness=2, color=PRIMARY))
    story.append(Spacer(1, 0.2 * inch))

    meta = [
        ["Repository", project.name],
        ["Analysis Date", datetime.utcnow().strftime("%B %d, %Y")],
        ["Total Files", str(project.total_files)],
        ["Migration Readiness", f"{project.migration_readiness_score:.1f}/100"],
    ]
    t = Table(meta, colWidths=[2 * inch, 4 * inch])
    t.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("TEXTCOLOR", (0, 0), (0, -1), PRIMARY),
        ("TEXTCOLOR", (1, 0), (1, -1), HexColor("#1e293b")),
        ("ROWBACKGROUNDS", (0, 0), (-1, -1), [LIGHT_GRAY, white]),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(t)
    story.append(PageBreak())

    # ── Score Cards ──────────────────────────────────────────────────────────
    story.append(Paragraph("Executive Summary", h2_style))
    story.append(HRFlowable(width="100%", thickness=1, color=HexColor("#e2e8f0")))
    story.append(Spacer(1, 0.1 * inch))
    story.append(Paragraph(executive_summary, body_style))
    story.append(Spacer(1, 0.2 * inch))

    score_data = [
        ["Metric", "Score", "Rating"],
        ["Migration Readiness", f"{project.migration_readiness_score:.1f}/100",
         "✓ Ready" if project.migration_readiness_score >= 70 else "⚠ Moderate" if project.migration_readiness_score >= 40 else "✗ Not Ready"],
        ["Average Risk", f"{project.avg_risk_score:.1f}/100",
         "Low" if project.avg_risk_score <= 30 else "Medium" if project.avg_risk_score <= 60 else "High" if project.avg_risk_score <= 80 else "Critical"],
        ["Technical Debt", f"{project.overall_debt_score:.1f}/100",
         "Low" if project.overall_debt_score <= 30 else "Medium" if project.overall_debt_score <= 60 else "High"],
        ["Security Score", f"{project.overall_security_score:.1f}/100",
         "Secure" if project.overall_security_score >= 80 else "Moderate" if project.overall_security_score >= 60 else "At Risk"],
    ]
    st = Table(score_data, colWidths=[2.5 * inch, 1.5 * inch, 2.5 * inch])
    st.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), DARK_BG),
        ("TEXTCOLOR", (0, 0), (-1, 0), white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [LIGHT_GRAY, white]),
        ("ALIGN", (1, 0), (-1, -1), "CENTER"),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#e2e8f0")),
    ]))
    story.append(st)
    story.append(PageBreak())

    # ── Top Risk Files ───────────────────────────────────────────────────────
    story.append(Paragraph("Top Risk Files", h2_style))
    story.append(HRFlowable(width="100%", thickness=1, color=HexColor("#e2e8f0")))
    story.append(Spacer(1, 0.1 * inch))

    top_files = sorted(files, key=lambda f: f.risk_score, reverse=True)[:15]
    risk_data = [["File", "Score", "Level", "LOC", "Complexity"]]
    for f in top_files:
        level = f.risk_level.value if f.risk_level else "Low"
        risk_data.append([
            f.relative_path[-50:] if len(f.relative_path) > 50 else f.relative_path,
            f"{f.risk_score:.1f}",
            level,
            str(f.loc),
            f"{f.cyclomatic_complexity:.1f}",
        ])

    rt = Table(risk_data, colWidths=[3 * inch, 0.8 * inch, 0.9 * inch, 0.7 * inch, 1.1 * inch])
    rt.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), DARK_BG),
        ("TEXTCOLOR", (0, 0), (-1, 0), white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [LIGHT_GRAY, white]),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#e2e8f0")),
    ]))
    story.append(rt)
    story.append(PageBreak())

    # ── Migration Roadmap ────────────────────────────────────────────────────
    story.append(Paragraph("Migration Roadmap", h2_style))
    story.append(HRFlowable(width="100%", thickness=1, color=HexColor("#e2e8f0")))
    story.append(Spacer(1, 0.1 * inch))

    for phase in phases:
        story.append(Paragraph(
            f"<b>Phase {phase['phase_number']}: {phase['name']}</b>",
            ParagraphStyle("PhaseTitle", parent=styles["Normal"],
                           fontSize=12, textColor=PRIMARY, spaceBefore=14, spaceAfter=4,
                           fontName="Helvetica-Bold")
        ))
        story.append(Paragraph(phase["description"], body_style))
        story.append(Paragraph(
            f"Files ({len(phase['files'])}): " + ", ".join(phase["files"][:8]) + ("..." if len(phase["files"]) > 8 else ""),
            ParagraphStyle("Files", parent=styles["Normal"], fontSize=8, textColor=MID_GRAY, spaceAfter=6)
        ))

    story.append(PageBreak())

    # ── Security Findings ────────────────────────────────────────────────────
    story.append(Paragraph("Security Analysis", h2_style))
    story.append(HRFlowable(width="100%", thickness=1, color=HexColor("#e2e8f0")))
    story.append(Spacer(1, 0.1 * inch))
    story.append(Paragraph(f"Overall Security Score: {project.overall_security_score:.1f}/100", body_style))

    sec_files = [f for f in files if f.security_findings]
    if sec_files:
        sec_data = [["File", "Finding Type", "Severity", "Line"]]
        for f in sec_files:
            for finding in f.security_findings[:3]:
                sec_data.append([
                    f.relative_path[-40:],
                    finding.finding_type.replace("_", " ").title(),
                    finding.severity.upper(),
                    str(finding.line_number or "-"),
                ])
        sft = Table(sec_data, colWidths=[2.8 * inch, 1.8 * inch, 1.2 * inch, 0.7 * inch])
        sft.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), DARK_BG),
            ("TEXTCOLOR", (0, 0), (-1, 0), white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [LIGHT_GRAY, white]),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#e2e8f0")),
        ]))
        story.append(sft)
    else:
        story.append(Paragraph("No critical security findings detected.", body_style))

    # ── Footer ───────────────────────────────────────────────────────────────
    story.append(Spacer(1, 0.5 * inch))
    story.append(HRFlowable(width="100%", thickness=1, color=HexColor("#e2e8f0")))
    story.append(Paragraph(
        f"Generated by LegacyLens on {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}",
        ParagraphStyle("Footer", parent=styles["Normal"], fontSize=8, textColor=MID_GRAY, alignment=TA_CENTER)
    ))

    doc.build(story)
    return buffer.getvalue()


def generate_docx(project: Project, files: List[SourceFile], phases: List[Dict], executive_summary: str) -> bytes:
    """Generate a DOCX executive report using python-docx."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Styles
    def set_color(run, hex_color: str):
        r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)

    # Title
    title = doc.add_heading("LegacyLens Migration Risk Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Repository: {project.name}")
    doc.add_paragraph(f"Date: {datetime.utcnow().strftime('%B %d, %Y')}")
    doc.add_paragraph(f"Migration Readiness: {project.migration_readiness_score:.1f}/100")
    doc.add_page_break()

    # Executive Summary
    doc.add_heading("Executive Summary", 1)
    doc.add_paragraph(executive_summary)

    # Score Cards
    doc.add_heading("Key Metrics", 2)
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    metrics = [
        ("Migration Readiness Score", f"{project.migration_readiness_score:.1f}/100"),
        ("Average Risk Score", f"{project.avg_risk_score:.1f}/100"),
        ("Technical Debt Score", f"{project.overall_debt_score:.1f}/100"),
        ("Security Score", f"{project.overall_security_score:.1f}/100"),
        ("Total Files Analyzed", str(project.total_files)),
    ]
    for i, (k, v) in enumerate(metrics):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    doc.add_page_break()

    # Top Risk Files
    doc.add_heading("Top Risk Files", 1)
    top_files = sorted(files, key=lambda f: f.risk_score, reverse=True)[:10]
    rt = doc.add_table(rows=1 + len(top_files), cols=4)
    rt.style = "Table Grid"
    headers = ["File", "Risk Score", "Risk Level", "LOC"]
    for j, h in enumerate(headers):
        rt.rows[0].cells[j].text = h
    for i, f in enumerate(top_files, 1):
        rt.rows[i].cells[0].text = f.relative_path
        rt.rows[i].cells[1].text = f"{f.risk_score:.1f}"
        rt.rows[i].cells[2].text = f.risk_level.value if f.risk_level else "Low"
        rt.rows[i].cells[3].text = str(f.loc)
    doc.add_page_break()

    # Migration Roadmap
    doc.add_heading("Migration Roadmap", 1)
    for phase in phases:
        doc.add_heading(f"Phase {phase['phase_number']}: {phase['name']}", 2)
        doc.add_paragraph(phase["description"])
        doc.add_paragraph(f"Files ({len(phase['files'])}): " + ", ".join(phase["files"][:10]))
    doc.add_page_break()

    # Security
    doc.add_heading("Security Analysis", 1)
    doc.add_paragraph(f"Overall Security Score: {project.overall_security_score:.1f}/100")
    sec_files = [f for f in files if f.security_findings]
    if sec_files:
        st = doc.add_table(rows=1, cols=4)
        st.style = "Table Grid"
        for j, h in enumerate(["File", "Type", "Severity", "Line"]):
            st.rows[0].cells[j].text = h
        for f in sec_files:
            for finding in f.security_findings[:5]:
                row = st.add_row()
                row.cells[0].text = f.relative_path
                row.cells[1].text = finding.finding_type
                row.cells[2].text = finding.severity
                row.cells[3].text = str(finding.line_number or "-")
    else:
        doc.add_paragraph("No critical security findings detected.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
